"""
Feature 4 — File ingestion (PDF / DOCX / Markdown / plain text).

Contract under test:

  * POST /ingest/file with a supported extension returns 201 plus
    {episode_id, file_name, file_type, file_sha256, char_count,
     chunk_count, memory_ids[]}.
  * Exactly ONE Episode is created per upload — the verbatim file text.
  * N Memories are created (one per chunk) and each carries
    source_episode_ids == [episode_id].
  * Search across the namespace surfaces ingested content.
  * Reverse traversal GET /episodes/{id}/memories returns every chunk
    derived from the file.
  * Unsupported extensions are rejected with HTTP 415.
  * Empty uploads are rejected with HTTP 400.

Tested formats: .md, .txt (always), .pdf (when reportlab is available),
.docx (when python-docx is available on the test runner — separate from
the server's python-docx which is always present after Dockerfile build).
"""
from __future__ import annotations

import io

import pytest


# ---------------------------------------------------------------------------
# Sample-file factories — keep them small and deterministic
# ---------------------------------------------------------------------------

MD_BODY = (
    "# Architecture Decision Record: PostgreSQL adoption\n\n"
    "We evaluated MongoDB, DynamoDB, and PostgreSQL for our analytics "
    "workload. PostgreSQL was selected because its mature SQL support "
    "and broad tooling ecosystem outweigh the JSON-native ergonomics of "
    "MongoDB for our use case. Compliance auditors require row-level "
    "access control, which PostgreSQL provides natively. The team's "
    "existing relational expertise reduces ramp-up risk. p99 latency "
    "must remain under 200 milliseconds for analytical reads — "
    "PostgreSQL benchmarks comfortably meet this on a managed service."
)


def _post_file(http, ns: str, filename: str, content: bytes, mime: str = "application/octet-stream"):
    """Upload one file. Returns the httpx.Response so callers can assert codes."""
    return http.post(
        "/api/v1/ingest/file",
        data={"namespace": ns},
        files={"file": (filename, io.BytesIO(content), mime)},
    )


# ---------------------------------------------------------------------------
# Happy-path uploads
# ---------------------------------------------------------------------------

class TestMarkdownIngest:
    def test_markdown_produces_episode_and_chunks(self, http, ns_episode: str):
        r = _post_file(http, ns_episode, "decision.md", MD_BODY.encode("utf-8"), "text/markdown")
        assert r.status_code == 201, f"{r.status_code} {r.text}"
        body = r.json()

        for key in ("episode_id", "file_name", "file_type", "file_sha256",
                    "char_count", "chunk_count", "memory_ids"):
            assert key in body, f"response missing {key}"

        assert body["file_type"] == "markdown"
        assert body["file_name"] == "decision.md"
        assert body["char_count"] == len(MD_BODY)
        assert body["chunk_count"] >= 1
        assert len(body["memory_ids"]) == body["chunk_count"]

    def test_chunks_link_back_to_episode(self, http, ns_episode: str):
        r = _post_file(http, ns_episode, "linkage.md", MD_BODY.encode("utf-8"), "text/markdown")
        assert r.status_code == 201
        body = r.json()
        ep_id = body["episode_id"]

        derived = http.get(
            f"/api/v1/episodes/{ep_id}/memories", params={"ns": ns_episode}
        ).json()
        derived_ids = {d["id"] for d in derived}
        assert set(body["memory_ids"]) == derived_ids, (
            "memory_ids in ingest response don't match /episodes/{id}/memories"
        )
        # Every derived chunk must reference the originating Episode.
        for d in derived:
            assert ep_id in d["source_episode_ids"], (
                f"derived chunk {d['id']} missing lineage to episode {ep_id}"
            )


class TestPlainTextIngest:
    def test_short_text_is_one_chunk(self, http, ns_episode: str):
        body = b"The quick brown fox jumps over the lazy dog."
        r = _post_file(http, ns_episode, "short.txt", body, "text/plain")
        assert r.status_code == 201
        data = r.json()
        assert data["file_type"] == "text"
        assert data["chunk_count"] == 1
        assert data["char_count"] == len(body.decode())


class TestPdfIngest:
    def test_pdf_extraction(self, http, ns_episode: str):
        pytest.importorskip("reportlab", reason="reportlab not installed on test runner")
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        for i, line in enumerate([
            "Acme Corporation Policy Memo",
            "Effective immediately, all production deploys require",
            "two-person review and a documented rollback plan.",
        ]):
            c.drawString(72, 720 - i * 20, line)
        c.showPage()
        c.save()
        pdf_bytes = buf.getvalue()

        r = _post_file(http, ns_episode, "policy.pdf", pdf_bytes, "application/pdf")
        assert r.status_code == 201, f"{r.status_code} {r.text}"
        data = r.json()
        assert data["file_type"] == "pdf"
        assert data["chunk_count"] >= 1
        # PDF text extraction must surface at least one of the original lines.
        ep = http.get(
            f"/api/v1/episodes/{data['episode_id']}", params={"ns": ns_episode}
        ).json()
        assert "two-person review" in ep["content"], (
            f"PDF text extraction failed; episode content: {ep['content'][:200]!r}"
        )


class TestDocxIngest:
    def test_docx_extraction(self, http, ns_episode: str):
        pytest.importorskip("docx", reason="python-docx not installed on test runner")
        from docx import Document

        doc = Document()
        doc.add_paragraph("Quarterly compliance review summary.")
        doc.add_paragraph(
            "All audit findings from Q3 have been remediated. "
            "No open critical issues remain at this time."
        )
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        r = _post_file(
            http, ns_episode, "review.docx", docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert r.status_code == 201, f"{r.status_code} {r.text}"
        data = r.json()
        assert data["file_type"] == "docx"

        ep = http.get(
            f"/api/v1/episodes/{data['episode_id']}", params={"ns": ns_episode}
        ).json()
        assert "audit findings" in ep["content"].lower()


# ---------------------------------------------------------------------------
# Search integration — ingested content is retrievable
# ---------------------------------------------------------------------------

class TestIngestRetrieval:
    def test_search_finds_ingested_content(self, http, ns_episode: str):
        r = _post_file(http, ns_episode, "retrievable.md", MD_BODY.encode(), "text/markdown")
        assert r.status_code == 201
        body = r.json()

        results = http.get(
            "/api/v1/memory/search",
            params={
                "q": "which database did we adopt for analytics",
                "ns": ns_episode,
                "mode": "vector",
                "top_k": 3,
            },
        ).json()
        assert results, "vector search returned nothing for ingested content"
        # At least one returned memory must come from the file we just ingested
        ingested = set(body["memory_ids"])
        assert ingested & {r["id"] for r in results}, (
            "search did not surface any chunk from the just-ingested file"
        )


# ---------------------------------------------------------------------------
# Error paths
# ---------------------------------------------------------------------------

class TestIngestErrors:
    def test_unsupported_extension_returns_415(self, http, ns_episode: str):
        r = _post_file(http, ns_episode, "binary.bin", b"\x00\x01\x02", "application/octet-stream")
        assert r.status_code == 415, f"expected 415, got {r.status_code}: {r.text}"

    def test_empty_upload_returns_400(self, http, ns_episode: str):
        r = _post_file(http, ns_episode, "empty.md", b"", "text/markdown")
        assert r.status_code == 400, f"expected 400, got {r.status_code}: {r.text}"

    def test_no_extractable_text_returns_400(self, http, ns_episode: str):
        # Whitespace-only file: passes the empty check but has no real content.
        r = _post_file(http, ns_episode, "blank.md", b"   \n\n  \t", "text/markdown")
        assert r.status_code == 400, f"expected 400, got {r.status_code}: {r.text}"
